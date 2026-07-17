"""Builds the final rewritten resume as a DOCX file."""
import io
from docx import Document
from docx.shared import Pt


def build_resume_docx(resume_text: str, accepted_rows: list[dict]) -> bytes:
    """Very simple MVP generator: takes the original resume text, applies
    accepted original->rewritten line swaps, and writes the result as a
    clean single-column DOCX (ATS-safe: no tables, no text boxes, no columns).
    """
    replacements = {row["original"].strip(): row["rewritten"].strip() for row in accepted_rows}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in resume_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        final_line = replacements.get(line, line)
        doc.add_paragraph(final_line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
