"""Extract plain text from an uploaded resume (PDF or DOCX)."""
import io
import pdfplumber
from docx import Document


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    raise ValueError("Unsupported file type — upload a .pdf or .docx resume.")


def _extract_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text_parts.append(page.extract_text() or "")
    return "\n".join(text_parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs).strip()


def detect_formatting_flags(filename: str, file_bytes: bytes) -> list[str]:
    """Heuristic checks for ATS-unfriendly formatting."""
    flags = []
    if filename.lower().endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        if doc.tables:
            flags.append("Table(s) detected — some ATS parsers skip or scramble table content.")
        section_count = sum(1 for s in doc.sections if s)
        if section_count > 1:
            flags.append("Multiple document sections detected — check for column layouts.")
    if filename.lower().endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                images = page.images
                if images:
                    flags.append("Embedded image(s) detected — text inside images is invisible to most ATS.")
                    break
    return flags
